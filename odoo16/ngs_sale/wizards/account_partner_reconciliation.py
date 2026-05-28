# -*- coding: utf-8 -*-
from odoo import api, fields, models, _
from odoo.exceptions import UserError, ValidationError
from datetime import datetime, timedelta
from dateutil.relativedelta import relativedelta
import base64
import io
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL


class AccountPartnerReconciliation(models.TransientModel):
    _name = 'account.partner.reconciliation'
    _description = 'Biên bản đối chiếu công nợ'

    company_id = fields.Many2one('res.company', string='Công ty', required=True, default=lambda self: self.env.company)
    partner_id = fields.Many2one('res.partner', string='Đối tác', required=True)
    date_from = fields.Date(string='Từ ngày', required=True, default=lambda self: fields.Date.today().replace(day=1))
    date_to = fields.Date(string='Đến ngày', required=True, 
                         default=lambda self: fields.Date.today().replace(day=1) + relativedelta(months=1, days=-1))
    initial_balance = fields.Monetary(string='Dư nợ đầu kỳ', currency_field='company_currency_id', readonly=True, compute='_compute_amounts')
    debit_amount = fields.Monetary(string='Phát sinh nợ trong kỳ', currency_field='company_currency_id', readonly=True, compute='_compute_amounts')
    credit_amount = fields.Monetary(string='Phát sinh có trong kỳ', currency_field='company_currency_id', readonly=True, compute='_compute_amounts')
    ending_balance = fields.Monetary(string='Dư nợ cuối kỳ', currency_field='company_currency_id', readonly=True, compute='_compute_amounts')
    company_currency_id = fields.Many2one(related='company_id.currency_id', string='Đơn vị tiền tệ')
    move_line_ids = fields.Many2many('account.move.line', string='Bút toán', compute='_compute_move_lines')
    
    # Fields for Word export
    word_file = fields.Binary('Word File', readonly=True, attachment=False)
    word_filename = fields.Char('Word Filename', readonly=True)
    
    @api.depends('partner_id', 'date_from', 'date_to', 'company_id')
    def _compute_move_lines(self):
        for rec in self:
            if not rec.partner_id or not rec.date_from or not rec.date_to:
                rec.move_line_ids = False
                continue
                
            domain = [
                ('partner_id', '=', rec.partner_id.id),
                ('company_id', '=', rec.company_id.id),
                ('date', '>=', rec.date_from),
                ('date', '<=', rec.date_to),
                ('parent_state', '=', 'posted'),
            ]
            
            rec.move_line_ids = self.env['account.move.line'].search(domain)
    
    @api.depends('partner_id', 'date_from', 'date_to', 'company_id', 'move_line_ids')
    def _compute_amounts(self):
        """
        Tính toán các chỉ tiêu cho biên bản đối chiếu công nợ:
        - Dư nợ đầu kỳ: Số dư trước ngày bắt đầu kỳ
        - Phát sinh nợ trong kỳ: Tổng phát sinh nợ TRONG khoảng thời gian được chọn (date_from đến date_to)
        - Phát sinh có trong kỳ: Tổng phát sinh có TRONG khoảng thời gian được chọn (date_from đến date_to)
        - Dư nợ cuối kỳ: Dư nợ đầu kỳ + Phát sinh nợ trong kỳ - Phát sinh có trong kỳ
        """
        for rec in self:
            if not rec.partner_id or not rec.date_from or not rec.date_to:
                rec.initial_balance = 0
                rec.debit_amount = 0
                rec.credit_amount = 0
                rec.ending_balance = 0
                continue

            # Sử dụng partner ledger report handler để tính toán chính xác
            report = self.env.ref('account_reports.partner_ledger_report')
            
            # Tạo options cho báo cáo với khoảng thời gian được chọn
            options = report._get_options({
                'date': {
                    'date_from': rec.date_from,
                    'date_to': rec.date_to,
                    'mode': 'range'
                },
                'company_ids': [rec.company_id.id],
                'unfold_all': True,
            })
            
            handler = self.env['account.partner.ledger.report.handler']
            
            # 1. Tính số dư đầu kỳ (trước ngày date_from)
            initial_balance_data = handler._get_initial_balance_values([rec.partner_id.id], options)
            first_column_group = next(iter(options['column_groups']))
            initial_balance_info = initial_balance_data.get(rec.partner_id.id, {}).get(first_column_group, {})
            rec.initial_balance = initial_balance_info.get('balance', 0.0)
            
            # 2. Tính phát sinh nợ/có TRONG KỲ (chỉ trong khoảng thời gian date_from đến date_to)
            # Tính trực tiếp từ move_line_ids đã được filter theo date_from và date_to
            # QUAN TRỌNG: 
            # - Phát sinh nợ: Chỉ tính từ INVOICES (hóa đơn) - move_type = 'out_invoice' (KH) hoặc 'in_invoice' (NCC)
            # - Phát sinh có: Chỉ tính từ PAYMENTS (thanh toán) - move_type = 'entry' và là payment
            # - Chỉ tính các move lines có tài khoản công nợ (receivable/payable)
            if rec.move_line_ids:
                # Filter chỉ lấy các move lines có tài khoản công nợ (receivable/payable)
                receivable_payable_lines = rec.move_line_ids.filtered(
                    lambda l: l.account_id.account_type in ['asset_receivable', 'liability_payable']
                )
                
                # Phát sinh nợ: Chỉ tính từ invoices (hóa đơn bán hàng cho KH, hóa đơn mua hàng cho NCC)
                # Với KH: move_type = 'out_invoice'
                # Với NCC: move_type = 'in_invoice'
                # Tính TẤT CẢ invoices phát sinh trong kỳ (theo logic nghiệp vụ: phát sinh = tất cả invoices trong kỳ)
                invoice_moves = receivable_payable_lines.filtered(
                    lambda l: l.move_id.move_type in ['out_invoice', 'in_invoice']
                )
                rec.debit_amount = sum(invoice_moves.mapped('debit'))
                
                # Phát sinh có: Chỉ tính từ payments (thanh toán)
                # Payment moves thường có move_type = 'entry' và có payment_id
                # Hoặc có thể check qua matched_debit_ids/matched_credit_ids (đã reconcile với invoice)
                payment_moves = receivable_payable_lines.filtered(
                    lambda l: l.move_id.move_type == 'entry' and l.move_id.payment_id
                )
                rec.credit_amount = sum(payment_moves.mapped('credit'))
            else:
                rec.debit_amount = 0.0
                rec.credit_amount = 0.0
            
            # 3. Tính dư nợ cuối kỳ theo công thức:
            # Dư nợ cuối kỳ = Dư nợ đầu kỳ + Phát sinh nợ trong kỳ - Phát sinh có trong kỳ
            rec.ending_balance = rec.initial_balance + rec.debit_amount - rec.credit_amount
    
    # def action_print_report(self):
    #     """Print the report in PDF format."""
    #     self.ensure_one()
    #     return self.env.ref('ngs_sale.action_report_account_partner_reconciliation').report_action(self)
    
    def _get_amount_text(self, amount):
        """Convert amount to text in Vietnamese."""
        return self.company_id.currency_id.amount_to_text(amount)
    
    def action_export_word(self):
        """Export the reconciliation statement as Word document."""
        self.ensure_one()
        
        # Create a new document
        doc = Document()
        
        # Set document properties
        doc.styles['Normal'].font.name = 'Times New Roman'
        doc.styles['Normal'].font.size = Pt(10)

        # Header: dùng bảng 1 dòng 2 cột, cột phải dài gấp 2 lần cột trái
        header_table = doc.add_table(rows=1, cols=2)
        header_table.autofit = False
        header_table.columns[0].width = Inches(1.5)
        header_table.columns[1].width = Inches(5.0)
        # Bên trái
        cell_left = header_table.cell(0, 0)
        cell_left_par = cell_left.paragraphs[0]
        cell_left_par.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = cell_left_par.add_run(self.company_id.name or '')
        run.bold = True
        run = cell_left_par.add_run('\nSố: ....../BBDC')
        run.bold = True
        # Bên phải
        cell_right = header_table.cell(0, 1)
        cell_right_par = cell_right.paragraphs[0]
        cell_right_par.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = cell_right_par.add_run('CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM')
        run.bold = True
        run.font.size = Pt(9)  # hoặc 10
        cell_right_par.add_run('\n')
        run = cell_right_par.add_run('Độc lập - Tự do - Hạnh phúc')
        run.italic = True
        run.bold = True
        run.font.size = Pt(9)  # hoặc 10
        cell_right_par.add_run('\n----------------------')
        cell_right_par.add_run(f'\nTp. Hồ Chí Minh, ngày {self.date_to.strftime("%d")} tháng {self.date_to.strftime("%m")} năm {self.date_to.strftime("%Y")}')
        doc.add_paragraph()  # Dòng trống
        # Tiêu đề
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run('BIÊN BẢN ĐỐI CHIẾU CÔNG NỢ').bold = True
        p.runs[0].font.size = Pt(12)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        # Lý do
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run('(V/v: Mua bán vật tư)').italic = True
        doc.add_paragraph()
        # Căn cứ
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.add_run(f'Căn cứ vào Hợp đồng số ... giữa {self.company_id.name} và {self.partner_id.name} về việc mua bán vật tư phục vụ nhu cầu sản xuất ...')
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.add_run('Căn cứ vào tình hình thực hiện Hợp đồng của hai đơn vị;')
        doc.add_paragraph('Hai bên gồm:')
        # Thông tin BÊN A
        p = doc.add_paragraph()
        p.add_run('BÊN A: ').bold = True
        p.add_run(f'{(self.partner_id.name or "").upper()}').bold = True
        p.add_run(' (BÊN MUA)').bold = True
        doc.add_paragraph(f'Địa chỉ: ').add_run((self.partner_id.street or "") + (", " + self.partner_id.street2 if self.partner_id.street2 else "") + (", " + self.partner_id.city if self.partner_id.city else "") + (", " + self.partner_id.state_id.name if self.partner_id.state_id else "") + (", " + self.partner_id.country_id.name if self.partner_id.country_id else ""))
        doc.add_paragraph('Mã số thuế: ').add_run(self.partner_id.vat or "")
        contacts = self.partner_id.child_ids.filtered(lambda c: c.type == 'invoice')
        p = doc.add_paragraph()
        p.add_run('Đại diện: ').bold = True
        if contacts:
            p.add_run(contacts[0].name.upper()).bold = True
        else:
            p.add_run("")
        p.add_run('   Chức vụ: ').bold = True
        if contacts and contacts[0].function:
            p.add_run(contacts[0].function.upper()).bold = True
        else:
            p.add_run("")
        # Thông tin BÊN B
        p = doc.add_paragraph()
        p.add_run('BÊN B: ').bold = True
        p.add_run(f'{(self.company_id.name or "").upper()}').bold = True
        p.add_run(' (BÊN BÁN)').bold = True
        doc.add_paragraph('Điện thoại: ').add_run(self.company_id.phone or "")
        doc.add_paragraph(f'Địa chỉ: ').add_run((self.company_id.partner_id.street or "") + (", " + self.company_id.partner_id.street2 if self.company_id.partner_id.street2 else "") + (", " + self.company_id.partner_id.city if self.company_id.partner_id.city else "") + (", " + self.company_id.partner_id.state_id.name if self.company_id.partner_id.state_id else "") + (", " + self.company_id.partner_id.country_id.name if self.company_id.partner_id.country_id else ""))
        if self.company_id.partner_id.bank_ids:
            for bank in self.company_id.partner_id.bank_ids:
                doc.add_paragraph('Tài khoản: ').add_run(bank.acc_number + ' - ' + (bank.bank_name or ""))
        doc.add_paragraph('MST: ').add_run(self.company_id.vat or "")
        company_contacts = self.company_id.partner_id.child_ids.filtered(lambda c: c.type == 'invoice')
        p = doc.add_paragraph()
        p.add_run('Đại diện: ').bold = True
        if company_contacts:
            p.add_run(company_contacts[0].name.upper()).bold = True
        else:
            p.add_run("")
        p.add_run('   Chức vụ: ').bold = True
        if company_contacts and company_contacts[0].function:
            p.add_run(company_contacts[0].function.upper()).bold = True
        else:
            p.add_run("")
        doc.add_paragraph()
        # Dòng mô tả thời gian đối chiếu
        p = doc.add_paragraph()
        p.add_run('Cùng nhau đối chiếu công nợ mua bán vật tư từ ngày ')
        run = p.add_run(self.date_from.strftime('%d/%m/%Y'))
        run.bold = True
        p.add_run(' đến ngày ')
        run = p.add_run(self.date_to.strftime('%d/%m/%Y'))
        run.bold = True
        p.add_run(', cụ thể như sau:')
        # Phần công nợ
        p = doc.add_paragraph()
        p.add_run('I. DƯ NỢ ĐẦU KỲ: ').bold = True
        p.add_run(f'{self.initial_balance:,.0f} đồng ').bold = True
        p.add_run('(1)').bold = True
        p = doc.add_paragraph()
        p.add_run('II. PHÁT SINH NỢ TRONG KỲ: ').bold = True
        p.add_run(f'{self.debit_amount:,.0f} đồng ').bold = True
        p.add_run('(2)').bold = True
        p = doc.add_paragraph()
        p.add_run('III. PHÁT SINH CÓ TRONG KỲ: ').bold = True
        p.add_run(f'{self.credit_amount:,.0f} đồng ').bold = True
        p.add_run('(3)').bold = True
        p = doc.add_paragraph()
        p.add_run('IV. DƯ NỢ CUỐI KỲ: (1) + (2) - (3) = ').bold = True
        p.add_run(f'{self.ending_balance:,.0f} đồng').bold = True
        # Đoạn kết
        doc.add_paragraph()
        amount_text = self._get_amount_text(self.ending_balance)
        p = doc.add_paragraph()
        p.add_run('Vậy đến ngày ')
        run = p.add_run(self.date_to.strftime('%d/%m/%Y'))
        run.bold = True
        p.add_run(f', {self.partner_id.name} còn nợ {self.company_id.name} số tiền là ')
        run = p.add_run(f'{self.ending_balance:,.0f} đồng')
        run.bold = True
        p.add_run(f' (Bằng chữ: {amount_text}).')
        p = doc.add_paragraph()
        p.add_run('Biên bản đối chiếu công nợ này được lập thành 02 (hai) bản, Bên A giữ 01 (một) bản có giá trị pháp lý như nhau.')
        # Chữ ký
        doc.add_paragraph()
        table = doc.add_table(rows=1, cols=2)
        table.style = 'TableGrid'
        table.cell(0, 0).paragraphs[0].add_run('ĐẠI DIỆN BÊN A').bold = True
        table.cell(0, 1).paragraphs[0].add_run('ĐẠI DIỆN BÊN B').bold = True
        # Set padding cho tất cả các paragraph
        for para in doc.paragraphs:
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(2)
        # Lưu file
        f = io.BytesIO()
        doc.save(f)
        f.seek(0)
        self.word_file = base64.b64encode(f.read())
        self.word_filename = f'Bien_ban_doi_chieu_cong_no_{self.partner_id.name}_{self.date_from.strftime("%d%m%Y")}_{self.date_to.strftime("%d%m%Y")}.docx'
        return {
            'name': _('Biên bản đối chiếu công nợ'),
            'type': 'ir.actions.act_window',
            'res_model': 'account.partner.reconciliation',
            'view_mode': 'form',
            'res_id': self.id,
            'target': 'new',
            'context': {'form_view_initial_mode': 'edit'},
        } 

    def get_partner_ledger_data(self, partner_id, date_from, date_to, company_id=None):
        """
        Lấy số dư đầu kỳ, phát sinh nợ/có, dư cuối kỳ và các dòng bút toán từ report handler sổ công nợ.
        """
        company = self.env['res.company'].browse(company_id) if company_id else self.env.company
        
        # Sử dụng partner ledger report handler
        report = self.env.ref('account_reports.partner_ledger_report')
        
        # Tạo options cho báo cáo
        options = report._get_options({
            'date': {
                'date_from': date_from,
                'date_to': date_to,
                'mode': 'range'
            },
            'company_ids': [company.id],
            'unfold_all': True,
        })
        
        handler = self.env['account.partner.ledger.report.handler']
        
        # Tính số dư đầu kỳ
        initial_balance_data = handler._get_initial_balance_values([partner_id], options)
        first_column_group = next(iter(options['column_groups']))
        initial_balance_info = initial_balance_data.get(partner_id, {}).get(first_column_group, {})
        initial_balance = initial_balance_info.get('balance', 0.0)
        
        # Tính phát sinh nợ/có trong kỳ
        query_results = handler._query_partners(options)
        partner_data = None
        for partner, results in query_results:
            if partner and partner.id == partner_id:
                partner_data = results.get(first_column_group, {})
                break
        
        if partner_data:
            debit_amount = partner_data.get('debit', 0.0)
            credit_amount = partner_data.get('credit', 0.0)
        else:
            debit_amount = 0.0
            credit_amount = 0.0
            
        # Tính dư nợ cuối kỳ
        ending_balance = initial_balance + debit_amount - credit_amount
        
        # Lấy các dòng bút toán
        aml_values = handler._get_aml_values(options, [partner_id])[partner_id]
        
        return {
            'initial_balance': initial_balance,
            'debit_amount': debit_amount,
            'credit_amount': credit_amount,
            'ending_balance': ending_balance,
            'move_lines': aml_values,
        } 